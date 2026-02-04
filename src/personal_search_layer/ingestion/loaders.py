"""Document loaders for ingestion."""

from __future__ import annotations

import csv
import hashlib
import json
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

from personal_search_layer.config import MAX_DOC_BYTES, MAX_PDF_PAGES
from personal_search_layer.models import LoadReport, LoadedDocument, TextBlock


SUPPORTED_SUFFIXES = {
    ".txt",
    ".md",
    ".pdf",
    ".html",
    ".htm",
    ".docx",
    ".ipynb",
    ".csv",
    ".tsv",
    ".json",
    ".geojson",
    ".r",
    ".py",
    ".sql",
    ".yml",
    ".yaml",
    ".sh",
}


def load_document(
    path: Path,
    *,
    max_doc_bytes: int = MAX_DOC_BYTES,
    max_pdf_pages: int = MAX_PDF_PAGES,
) -> tuple[LoadedDocument | None, LoadReport]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported file type: {path.suffix}")

    if suffix == ".pdf":
        source_type = "pdf"
    elif suffix in {".html", ".htm"}:
        source_type = "html"
    elif suffix == ".docx":
        source_type = "docx"
    elif suffix == ".ipynb":
        source_type = "notebook"
    elif suffix in {".csv", ".tsv"}:
        source_type = "csv"
    elif suffix in {".json", ".geojson"}:
        source_type = "json"
    else:
        source_type = "text"
    bytes_total = path.stat().st_size
    if bytes_total > max_doc_bytes:
        report = LoadReport(
            source_path=str(path),
            source_type=source_type,
            bytes_total=bytes_total,
            pages_total=None,
            pages_loaded=0,
            pages_skipped_empty=0,
            pages_skipped_limit=0,
            skip_reason="file_too_large",
        )
        return None, report

    if suffix == ".pdf":
        blocks, report = _load_pdf(path, max_pages=max_pdf_pages)
    elif suffix in {".html", ".htm"}:
        blocks = [_load_html(path)]
        report = LoadReport(
            source_path=str(path),
            source_type=source_type,
            bytes_total=bytes_total,
            pages_total=1,
            pages_loaded=1,
            pages_skipped_empty=0,
            pages_skipped_limit=0,
            skip_reason=None,
        )
    elif suffix == ".docx":
        try:
            blocks = [_load_docx(path)]
            report = LoadReport(
                source_path=str(path),
                source_type=source_type,
                bytes_total=bytes_total,
                pages_total=1,
                pages_loaded=1,
                pages_skipped_empty=0,
                pages_skipped_limit=0,
                skip_reason=None,
            )
        except Exception:
            report = LoadReport(
                source_path=str(path),
                source_type=source_type,
                bytes_total=bytes_total,
                pages_total=None,
                pages_loaded=0,
                pages_skipped_empty=0,
                pages_skipped_limit=0,
                skip_reason="docx_parse_error",
            )
            return None, report
    elif suffix == ".ipynb":
        blocks = [_load_ipynb(path)]
        report = LoadReport(
            source_path=str(path),
            source_type=source_type,
            bytes_total=bytes_total,
            pages_total=1,
            pages_loaded=1,
            pages_skipped_empty=0,
            pages_skipped_limit=0,
            skip_reason=None,
        )
    elif suffix in {".csv", ".tsv"}:
        delimiter = "," if suffix == ".csv" else "\t"
        blocks = [_load_csv(path, delimiter=delimiter)]
        report = LoadReport(
            source_path=str(path),
            source_type=source_type,
            bytes_total=bytes_total,
            pages_total=1,
            pages_loaded=1,
            pages_skipped_empty=0,
            pages_skipped_limit=0,
            skip_reason=None,
        )
    elif suffix in {".json", ".geojson"}:
        blocks = [_load_json(path)]
        report = LoadReport(
            source_path=str(path),
            source_type=source_type,
            bytes_total=bytes_total,
            pages_total=1,
            pages_loaded=1,
            pages_skipped_empty=0,
            pages_skipped_limit=0,
            skip_reason=None,
        )
    else:
        blocks = [_load_text(path)]
        report = LoadReport(
            source_path=str(path),
            source_type=source_type,
            bytes_total=bytes_total,
            pages_total=1,
            pages_loaded=1,
            pages_skipped_empty=0,
            pages_skipped_limit=0,
            skip_reason=None,
        )

    content_hash = _hash_blocks(blocks)
    return (
        LoadedDocument(
            source_path=str(path),
            source_type=source_type,
            title=path.stem,
            blocks=blocks,
            content_hash=content_hash,
        ),
        report,
    )


def _load_pdf(
    path: Path, *, max_pages: int = MAX_PDF_PAGES
) -> tuple[list[TextBlock], LoadReport]:
    source_type = "pdf"
    bytes_total = path.stat().st_size
    pages_loaded = 0
    pages_skipped_empty = 0
    pages_skipped_limit = 0
    try:
        reader = PdfReader(str(path))
    except Exception:
        report = LoadReport(
            source_path=str(path),
            source_type=source_type,
            bytes_total=bytes_total,
            pages_total=None,
            pages_loaded=0,
            pages_skipped_empty=0,
            pages_skipped_limit=0,
            skip_reason="pdf_parse_error",
        )
        return [], report

    total_pages = len(reader.pages)
    blocks: list[TextBlock] = []
    for idx, page in enumerate(reader.pages, start=1):
        if idx > max_pages:
            pages_skipped_limit = total_pages - max_pages
            break
        text = page.extract_text() or ""
        if text.strip():
            blocks.append(TextBlock(text=text, page=idx))
            pages_loaded += 1
        else:
            pages_skipped_empty += 1
    report = LoadReport(
        source_path=str(path),
        source_type=source_type,
        bytes_total=bytes_total,
        pages_total=total_pages,
        pages_loaded=pages_loaded,
        pages_skipped_empty=pages_skipped_empty,
        pages_skipped_limit=pages_skipped_limit,
        skip_reason=None,
    )
    return blocks, report


def _load_text(path: Path) -> TextBlock:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return TextBlock(text=text)


def _load_docx(path: Path) -> TextBlock:
    doc = Document(str(path))
    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    return TextBlock(text="\n".join(paragraphs))


def _load_ipynb(path: Path) -> TextBlock:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    try:
        payload = json.loads(raw)
    except json.JSONDecodeError:
        return TextBlock(text=raw)
    cells = payload.get("cells", [])
    parts: list[str] = []
    for cell in cells:
        cell_type = cell.get("cell_type", "cell")
        source = cell.get("source", [])
        content = "".join(source).strip()
        if not content:
            continue
        parts.append(f"[{cell_type}]\n{content}")
    text = "\n\n".join(parts) if parts else raw
    return TextBlock(text=text)


def _load_csv(path: Path, *, delimiter: str) -> TextBlock:
    rows: list[str] = []
    with path.open(encoding="utf-8", errors="ignore", newline="") as handle:
        reader = csv.reader(handle, delimiter=delimiter)
        for row in reader:
            rows.append("\t".join(row))
    return TextBlock(text="\n".join(rows))


def _load_json(path: Path) -> TextBlock:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    try:
        payload = json.loads(raw)
    except json.JSONDecodeError:
        return TextBlock(text=raw)
    text = json.dumps(payload, ensure_ascii=False, indent=2)
    return TextBlock(text=text)


def _load_html(path: Path) -> TextBlock:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(raw, "html.parser")
    text = soup.get_text(separator=" ", strip=True)
    return TextBlock(text=text)


def _hash_blocks(blocks: list[TextBlock]) -> str:
    content = "\n".join(block.text for block in blocks)
    return hashlib.sha256(content.encode("utf-8")).hexdigest()
